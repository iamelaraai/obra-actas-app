import io
import re
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

st.set_page_config(page_title="Actas de Obra - Generador", page_icon="📋", layout="wide")

ESTADOS = ["Cumplido", "En proceso", "No cumplido", "Cumplido parcialmente", "Pendiente por definir"]


def clean_text(v):
    if pd.isna(v):
        return ""
    return re.sub(r"\s+", " ", str(v)).strip()


def parse_fecha_estado(raw: str):
    txt = clean_text(raw)
    if not txt:
        return "", "", txt

    lines = [x.strip() for x in re.split(r"\n|\r", txt) if x.strip()]
    fecha, estado = "", ""

    if lines:
        m = re.search(r"\b(\d{1,2}/\d{1,2}/\d{2,4})\b", lines[0])
        if m:
            fecha = m.group(1)

    lowered = txt.lower()
    for e in ESTADOS:
        if e.lower() in lowered:
            estado = e
            break

    return fecha, estado, txt


def extract_block(df: pd.DataFrame, actor: str, c_comp: int, c_compo: int, c_resp: int, c_fc: int, c_obs: int):
    rows = []
    for _, row in df.iterrows():
        compromiso = clean_text(row.iloc[c_comp]) if c_comp < len(row) else ""
        componente = clean_text(row.iloc[c_compo]) if c_compo < len(row) else ""
        responsable = clean_text(row.iloc[c_resp]) if c_resp < len(row) else ""
        fecha_raw = clean_text(row.iloc[c_fc]) if c_fc < len(row) else ""
        observ = clean_text(row.iloc[c_obs]) if c_obs < len(row) else ""

        if not compromiso:
            continue

        fecha_limite, estado, fecha_comentarios = parse_fecha_estado(fecha_raw)

        rows.append(
            {
                "Actor": actor,
                "Compromiso": compromiso,
                "Componente": componente,
                "Responsable": responsable,
                "Fecha límite": fecha_limite,
                "Estado": estado,
                "Fecha/Comentarios (raw)": fecha_comentarios,
                "Observación seguimiento": observ,
            }
        )
    return rows


def parse_sheet(df: pd.DataFrame, acta_no: str, fecha_comite: str):
    blocks = []
    blocks += extract_block(df, "EDU", 1, 2, 3, 4, 5)
    blocks += extract_block(df, "Contratista", 7, 8, 9, 10, 11)
    blocks += extract_block(df, "Interventoría", 13, 14, 15, 16, 17)

    out = pd.DataFrame(blocks)
    if out.empty:
        return out

    out.insert(0, "Acta No", acta_no)
    out.insert(1, "Fecha comité", fecha_comite)
    out.insert(2, "ID compromiso", [f"A{acta_no}-{a[:3].upper()}-{str(i+1).zfill(3)}" for i, a in enumerate(out["Actor"].tolist())])
    return out


def semaforo_estado(estado: str):
    e = (estado or "").strip().lower()
    if e == "cumplido":
        return "🟢"
    if e in ["en proceso", "cumplido parcialmente", "pendiente por definir"]:
        return "🟡"
    if e == "no cumplido":
        return "🔴"
    return "⚪"


def build_summary(df: pd.DataFrame):
    total = len(df)
    by_estado = df["Estado"].fillna("").value_counts(dropna=False).to_dict()
    by_actor = df["Actor"].value_counts(dropna=False).to_dict()

    lines = [f"Total compromisos: {total}", "\nPor actor:"]
    for k, v in by_actor.items():
        lines.append(f"- {k}: {v}")

    lines.append("\nPor estado:")
    for k, v in by_estado.items():
        kk = k if k else "(sin estado)"
        lines.append(f"- {kk}: {v}")

    return "\n".join(lines)


def build_acta_text(df: pd.DataFrame):
    parts = ["## Compromisos, comentarios y observaciones\n"]
    for actor in ["EDU", "Contratista", "Interventoría"]:
        sub = df[df["Actor"] == actor]
        if sub.empty:
            continue
        parts.append(f"### {actor}")
        for _, r in sub.iterrows():
            estado = r.get("Estado", "") or "Sin estado"
            fecha = r.get("Fecha límite", "") or "Sin fecha"
            sem = semaforo_estado(estado)
            parts.append(
                f"- {sem} **{r['Compromiso']}** ({r['Componente']}) — Responsable: {r['Responsable']}. "
                f"Estado: {estado}. Fecha: {fecha}. Observación: {r.get('Observación seguimiento','') or 'N/A'}."
            )
        parts.append("")
    return "\n".join(parts)


def build_full_acta_md(meta: dict, df: pd.DataFrame, resumen_reunion: str):
    head = [
        f"# Acta de Comité de Obra No. {meta.get('acta_no','')}",
        f"**Proyecto:** {meta.get('proyecto','')}",
        f"**Fecha:** {meta.get('fecha','')}",
        f"**Lugar:** {meta.get('lugar','')}",
        f"**Hora inicio:** {meta.get('hora_inicio','')}",
        f"**Hora fin:** {meta.get('hora_fin','')}",
        "",
        "## 1) Resumen ejecutivo",
        meta.get('resumen_ejecutivo','') or "(Completar)",
        "",
        "## 2) Decisiones y temas relevantes",
        (resumen_reunion or "(Pegar aquí resumen generado desde transcripción)"),
        "",
        "## 3) Compromisos, comentarios y observaciones",
    ]
    return "\n".join(head) + "\n\n" + build_acta_text(df)


def to_docx_from_md(texto: str):
    doc = Document()
    for line in texto.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            doc.add_paragraph(line)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def build_transcript_prompt(transcript: str, contexto: str):
    return f"""Eres asistente de interventoría de obra. Redacta sección de acta en español formal y técnico.

Contexto del proyecto:
{contexto}

Transcripción de reunión:
{transcript}

Devuelve SOLO con esta estructura:
1) Decisiones tomadas
2) Avances reportados
3) Riesgos y atrasos críticos
4) Compromisos nuevos (tabla en markdown con: compromiso, responsable, fecha límite, componente)
5) Observaciones de cierre
"""


def generar_observacion_breve(estado: str, actor: str, compromiso: str, notas: str, estilo: str = "Interventoría formal"):
    est = (estado or "").strip().lower()
    actor_txt = actor or "El responsable"
    c = clean_text(compromiso)
    n = clean_text(notas)

    if estilo == "Ejecutivo corto":
        if est == "cumplido":
            return f"Compromiso cumplido por {actor_txt}. {n}".strip()
        if est == "no cumplido":
            return f"Compromiso no cumplido por {actor_txt}. {n} Se reprograma seguimiento.".strip()
        if est == "cumplido parcialmente":
            return f"Cumplimiento parcial por {actor_txt}. {n}".strip()
        if est == "en proceso":
            return f"Compromiso en proceso por {actor_txt}. {n}".strip()
        return f"Estado pendiente de confirmar. {n}".strip()

    if estilo == "Operativo campo":
        if est == "cumplido":
            return f"{actor_txt} reporta ejecución completa del compromiso: {c}. Evidencia/nota: {n or 'sin novedad'}"
        if est == "no cumplido":
            return f"{actor_txt} reporta no cumplimiento del compromiso: {c}. Causa/nota: {n or 'pendiente información'}. Acción: reprogramar y hacer seguimiento en próximo comité."
        if est == "cumplido parcialmente":
            return f"{actor_txt} reporta avance parcial del compromiso: {c}. Avance: {n or 'sin detalle'}. Acción: completar pendiente y validar cierre."
        if est == "en proceso":
            return f"{actor_txt} reporta compromiso en ejecución: {c}. Estado actual: {n or 'en curso'}. Acción: mantener seguimiento."
        return f"{actor_txt} reporta novedad sobre compromiso: {c}. Nota: {n or 'sin detalle'}."

    if estilo == "Neutro estándar":
        if est == "cumplido":
            return f"{actor_txt} informa cumplimiento del compromiso: {c}. {('Detalle: ' + n) if n else ''}".strip()
        if est == "no cumplido":
            return f"{actor_txt} informa no cumplimiento del compromiso: {c}. {('Detalle: ' + n) if n else ''} Se traslada seguimiento.".strip()
        if est == "cumplido parcialmente":
            return f"{actor_txt} informa cumplimiento parcial del compromiso: {c}. {('Detalle: ' + n) if n else ''}".strip()
        if est == "en proceso":
            return f"{actor_txt} informa que el compromiso está en proceso: {c}. {('Detalle: ' + n) if n else ''}".strip()
        return f"Observación sobre el compromiso '{c}': {n}" if n else f"Observación pendiente de completar para el compromiso: {c}."

    # Interventoría formal (default)
    if est == "cumplido":
        base = f"Desde {actor_txt} se informa que se dio cumplimiento al compromiso: {c}."
        extra = f" Se deja como soporte: {n}." if n else ""
        return base + extra

    if est == "no cumplido":
        base = f"Desde {actor_txt} se informa que no se ha dado cumplimiento al compromiso: {c}."
        extra = f" Se reporta: {n}." if n else ""
        return base + extra + " Se traslada para seguimiento en el próximo comité."

    if est == "cumplido parcialmente":
        base = f"Desde {actor_txt} se informa cumplimiento parcial del compromiso: {c}."
        extra = f" Avance reportado: {n}." if n else ""
        return base + extra + " Se requiere completar actividades pendientes."

    if est == "en proceso":
        base = f"Desde {actor_txt} se informa que el compromiso {c} se encuentra en proceso."
        extra = f" Se reporta: {n}." if n else ""
        return base + extra + " Se mantiene seguimiento en el próximo comité."

    return f"Observación sobre el compromiso '{c}': {n}" if n else f"Observación pendiente de completar para el compromiso: {c}."


def to_xlsx_bytes(df: pd.DataFrame, sheet_name: str = "compromisos"):
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    bio.seek(0)
    return bio


def to_official_template_bytes(template_file, df_simple: pd.DataFrame, sheet_name: str = "Acta 19"):
    """Rellena plantilla oficial sin romper formato y colorea FECHA/COMENTARIOS por estado."""
    wb = load_workbook(template_file)
    if sheet_name not in wb.sheetnames:
        sheet_name = wb.sheetnames[-1]
    ws = wb[sheet_name]

    # Limpia rango de datos (filas de compromisos) en columnas B:F, H:L, N:R
    for row in range(4, 220):
        for col in [2, 3, 4, 5, 6, 8, 9, 10, 11, 12, 14, 15, 16, 17, 18]:
            cell = ws.cell(row=row, column=col)
            cell.value = None
            # limpiamos relleno solo en columna fecha/comentarios de cada bloque
            if col in [5, 11, 17]:
                cell.fill = PatternFill(fill_type=None)

    # Mapeo por responsable/actor -> columna base
    base_cols = {
        "edu": 2,            # B
        "contratista": 8,    # H
        "interventoría": 14, # N
        "interventoria": 14,
    }

    # Colores de estado (como en tu ejemplo)
    estado_fill = {
        "cumplido": PatternFill(fill_type="solid", start_color="92D050", end_color="92D050"),
        "en proceso": PatternFill(fill_type="solid", start_color="FFFF00", end_color="FFFF00"),
        "cumplido parcialmente": PatternFill(fill_type="solid", start_color="FFFF00", end_color="FFFF00"),
        "pendiente por definir": PatternFill(fill_type="solid", start_color="FFFF00", end_color="FFFF00"),
        "no cumplido": PatternFill(fill_type="solid", start_color="FF0000", end_color="FF0000"),
    }

    # contador de fila por bloque para que cada columna tenga sus propios renglones
    next_row = {2: 4, 8: 4, 14: 4}

    for _, r in df_simple.iterrows():
        responsable = clean_text(r.get("Responsable", "")).lower()
        actor = clean_text(r.get("Actor", "")).lower()

        # prioridad: responsable; fallback: actor
        base = base_cols.get(responsable) or base_cols.get(actor)
        if not base:
            # si viene raro, se manda a EDU para no perder el dato
            base = 2

        row_i = next_row[base]
        next_row[base] += 1

        ws.cell(row=row_i, column=base).value = clean_text(r.get("Compromiso", ""))
        ws.cell(row=row_i, column=base + 1).value = clean_text(r.get("Componente", ""))
        ws.cell(row=row_i, column=base + 2).value = clean_text(r.get("Responsable", ""))

        fecha = clean_text(r.get("Fecha límite", ""))
        estado = clean_text(r.get("Estado", ""))
        fc = "\n".join([x for x in [fecha, estado] if x])

        fc_cell = ws.cell(row=row_i, column=base + 3)
        fc_cell.value = fc

        fill = estado_fill.get(estado.lower())
        if fill:
            fc_cell.fill = fill

        ws.cell(row=row_i, column=base + 4).value = clean_text(r.get("Observación seguimiento", ""))

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


st.title("📋 Generador de Actas de Obra (teach-friendly)")
st.caption("Excel → compromisos normalizados → sección/acta Word + apoyo de transcripción")

tab0, tab1, tab2, tab3 = st.tabs([
    "0) Captura guiada de compromisos",
    "1) Compromisos desde Excel",
    "2) Transcripción y resumen",
    "3) Acta completa (1 clic)",
])

editable = None
texto_acta = ""

with tab0:
    st.subheader("Ingreso simple de compromisos (sin pelear con Excel)")
    st.caption("Tu hermano llena esta tabla guiada y descarga un Excel limpio listo para usar en el Tab 1.")

    c0a, c0b, c0c = st.columns(3)
    with c0a:
        acta_no_form = st.text_input("Acta No (captura)", value="19", key="acta_no_form")
    with c0b:
        fecha_form = st.text_input("Fecha comité (captura)", value=datetime.now().strftime("%d/%m/%Y"), key="fecha_form")
    with c0c:
        actor_default = st.selectbox("Actor por defecto", ["EDU", "Contratista", "Interventoría"], key="actor_default")

    base_cols = [
        "Acta No",
        "Fecha comité",
        "Actor",
        "Compromiso",
        "Componente",
        "Responsable",
        "Fecha límite",
        "Estado",
        "Observación seguimiento",
    ]
    if "captura_df" not in st.session_state:
        st.session_state["captura_df"] = pd.DataFrame(
            [
                {
                    "Acta No": acta_no_form,
                    "Fecha comité": fecha_form,
                    "Actor": actor_default,
                    "Compromiso": "",
                    "Componente": "Técnico",
                    "Responsable": "",
                    "Fecha límite": "",
                    "Estado": "En proceso",
                    "Observación seguimiento": "",
                }
            ]
        )

    # sincroniza metadatos por defecto en filas vacías
    base_df = st.session_state["captura_df"].copy()
    for col in base_cols:
        if col not in base_df.columns:
            base_df[col] = ""
    base_df = base_df[base_cols]
    base_df.loc[base_df["Acta No"].astype(str).str.strip() == "", "Acta No"] = acta_no_form
    base_df.loc[base_df["Fecha comité"].astype(str).str.strip() == "", "Fecha comité"] = fecha_form

    edited_df = st.data_editor(
        base_df,
        use_container_width=True,
        num_rows="dynamic",
        key="captura_editor",
        column_config={
            "Actor": st.column_config.SelectboxColumn("Actor", options=["EDU", "Contratista", "Interventoría"]),
            "Estado": st.column_config.SelectboxColumn("Estado", options=[""] + ESTADOS),
        },
    )

    # Evita sobreescrituras accidentales cuando Streamlit re-renderiza en estados intermedios
    if isinstance(edited_df, pd.DataFrame) and set(base_cols).issubset(set(edited_df.columns)):
        if len(edited_df) > 0 or len(st.session_state.get("captura_df", pd.DataFrame())) == 0:
            st.session_state["captura_df"] = edited_df[base_cols].copy()

    # Tabla visible (sin filtrar) y tabla procesada (solo filas con compromiso)
    captura_df_visible = st.session_state["captura_df"].copy()
    captura_df = captura_df_visible[captura_df_visible["Compromiso"].astype(str).str.strip() != ""].copy()

    if not captura_df.empty:
        captura_df.insert(
            2,
            "ID compromiso",
            [f"A{a}-{str(act)[:3].upper()}-{str(i+1).zfill(3)}" for i, (a, act) in enumerate(zip(captura_df["Acta No"], captura_df["Actor"]))],
        )

        st.success(f"Compromisos capturados: {len(captura_df)}")

        st.markdown("### ✍️ Redactor breve de observaciones")
        st.caption("Escribe notas cortas y genera una observación formal lista para pegar en la columna de seguimiento.")
        rx1, rx2 = st.columns(2)
        with rx1:
            idx = st.selectbox("Selecciona compromiso", options=list(captura_df.index), format_func=lambda i: f"{captura_df.loc[i, 'Actor']} · {str(captura_df.loc[i, 'Compromiso'])[:80]}")
            est_sel = st.selectbox("Estado para redactar", options=ESTADOS, index=max(ESTADOS.index(captura_df.loc[idx, 'Estado']) if captura_df.loc[idx, 'Estado'] in ESTADOS else 1, 0))
        with rx2:
            estilo_sel = st.selectbox(
                "Estilo de redacción",
                ["Interventoría formal", "Ejecutivo corto", "Operativo campo", "Neutro estándar"],
                index=0,
            )
            notas_cortas = st.text_area("Notas rápidas (2-10 palabras)", placeholder="ej: pendiente respuesta de SIF sobre alcance de red", height=90)

        obs_generada = generar_observacion_breve(
            est_sel,
            str(captura_df.loc[idx, "Actor"]),
            str(captura_df.loc[idx, "Compromiso"]),
            notas_cortas,
            estilo_sel,
        )
        st.text_area("Observación sugerida", value=obs_generada, height=110)

        if st.button("Usar observación sugerida en la fila seleccionada"):
            full_df = st.session_state["captura_df"].copy()
            if idx in full_df.index:
                full_df.loc[idx, "Estado"] = est_sel
                full_df.loc[idx, "Observación seguimiento"] = obs_generada
                st.session_state["captura_df"] = full_df
                st.success("Observación aplicada en la fila seleccionada.")
                st.rerun()
        st.download_button(
            "⬇️ Descargar Excel de compromisos (simple)",
            data=to_xlsx_bytes(captura_df, sheet_name=f"Acta_{acta_no_form}"),
            file_name=f"compromisos_acta_{acta_no_form}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        st.download_button(
            "⬇️ Descargar CSV (simple)",
            data=captura_df.to_csv(index=False).encode("utf-8"),
            file_name=f"compromisos_acta_{acta_no_form}.csv",
            mime="text/csv",
        )

        st.markdown("#### 🧩 Exportar en formato oficial (igual al institucional)")
        st.caption("Sube la plantilla oficial (xlsx) y la app rellenará los compromisos manteniendo formato visual.")
        t1, t2 = st.columns([2, 1])
        with t1:
            plantilla = st.file_uploader("Plantilla oficial (.xlsx)", type=["xlsx"], key="plantilla_oficial")
        with t2:
            nombre_hoja = st.text_input("Pestaña destino", value=f"Acta {acta_no_form}", key="hoja_destino")

        if plantilla is not None:
            oficial_bytes = to_official_template_bytes(plantilla, captura_df, sheet_name=nombre_hoja)
            st.download_button(
                "⬇️ Descargar Excel en formato oficial",
                data=oficial_bytes,
                file_name=f"compromisos_acta_{acta_no_form}_formato_oficial.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        st.info("Tip: este Excel simple te sirve como base maestra. También puedes pasarlo directo a Tab 3 para generar acta completa.")
    else:
        st.warning("Agrega al menos un compromiso con texto para habilitar descargas.")

with tab1:
    archivo = st.file_uploader("Sube el Excel de compromisos", type=["xlsx", "xlsm", "xlsb"], key="excel")

    if archivo is not None:
        xls = pd.ExcelFile(archivo)
        sheet = st.selectbox("Pestaña", xls.sheet_names, index=max(len(xls.sheet_names) - 1, 0))

        c1, c2 = st.columns(2)
        with c1:
            acta_no = st.text_input("Acta No", value=re.sub(r"\D", "", sheet) or "")
        with c2:
            fecha_comite = st.text_input("Fecha comité", value=datetime.now().strftime("%d/%m/%Y"))

        # Soporta dos formatos:
        # A) Formato legado por bloques (EDU/Contratista/Interventoría)
        # B) Formato simple (tabla maestra con columnas)
        data = pd.DataFrame()

        try:
            simple = pd.read_excel(archivo, sheet_name=sheet)
            req = {"Acta No", "Fecha comité", "Actor", "Compromiso", "Componente", "Responsable", "Fecha límite", "Estado", "Observación seguimiento"}
            if req.issubset(set(simple.columns)):
                data = simple.copy()
                if "ID compromiso" not in data.columns:
                    data.insert(
                        2,
                        "ID compromiso",
                        [f"A{a}-{str(act)[:3].upper()}-{str(i+1).zfill(3)}" for i, (a, act) in enumerate(zip(data["Acta No"], data["Actor"]))],
                    )
        except Exception:
            pass

        if data.empty:
            raw = pd.read_excel(archivo, sheet_name=sheet, header=None)
            data = parse_sheet(raw, acta_no=acta_no or "", fecha_comite=fecha_comite)

        if data.empty:
            st.warning("No se detectaron compromisos con la estructura esperada en esta pestaña.")
        else:
            editable = st.data_editor(
                data,
                use_container_width=True,
                num_rows="dynamic",
                column_config={"Estado": st.column_config.SelectboxColumn("Estado", options=[""] + ESTADOS)},
                key="editor_compromisos"
            )
            texto_acta = build_acta_text(editable)

            st.text(build_summary(editable))
            st.text_area("Sección de compromisos (copiar/pegar en Word)", value=texto_acta, height=260)

            c3, c4, c5 = st.columns(3)
            with c3:
                st.download_button(
                    "⬇️ CSV normalizado",
                    data=editable.to_csv(index=False).encode("utf-8"),
                    file_name=f"compromisos_acta_{acta_no or sheet}.csv",
                    mime="text/csv",
                )
            with c4:
                st.download_button(
                    "⬇️ Markdown",
                    data=texto_acta.encode("utf-8"),
                    file_name=f"seccion_compromisos_acta_{acta_no or sheet}.md",
                    mime="text/markdown",
                )
            with c5:
                st.download_button(
                    "⬇️ Word sección",
                    data=to_docx_from_md(texto_acta),
                    file_name=f"seccion_compromisos_acta_{acta_no or sheet}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

with tab2:
    st.caption("Pega transcripción y obtén prompt maestro para Claude/Codex/ChatGPT.")

    contexto = st.text_area(
        "Contexto del proyecto",
        value="Proyecto: Parque Primavera Norte. Documento: acta de comité de obra.",
        height=80,
        key="contexto"
    )
    transcript = st.text_area("Transcripción de reunión", height=220, placeholder="Pega aquí la transcripción...", key="transcript")

    if transcript.strip():
        prompt = build_transcript_prompt(transcript.strip(), contexto.strip())
        st.code(prompt, language="markdown")
        st.download_button(
            "⬇️ Descargar prompt (.txt)",
            data=prompt.encode("utf-8"),
            file_name="prompt_resumen_acta.txt",
            mime="text/plain",
        )

with tab3:
    st.subheader("Generar acta completa de hoy")
    st.caption("Usa lo cargado en Tab 1 + un resumen de reunión (manual o generado con IA).")

    p1, p2 = st.columns(2)
    with p1:
        proyecto = st.text_input("Proyecto", value="Parque Primavera Norte")
        lugar = st.text_input("Lugar", value="Campamento de obra")
        hora_inicio = st.text_input("Hora inicio", value="09:30 AM")
    with p2:
        no_acta = st.text_input("No. Acta", value="")
        fecha = st.text_input("Fecha", value=datetime.now().strftime("%d/%m/%Y"))
        hora_fin = st.text_input("Hora fin", value="11:30 AM")

    resumen_ejecutivo = st.text_area("Resumen ejecutivo", height=100, placeholder="3-5 líneas del estado general...")
    resumen_reunion = st.text_area("Decisiones/avances/riesgos (desde transcripción)", height=180)

    if editable is not None and len(editable) > 0:
        md_full = build_full_acta_md(
            {
                "acta_no": no_acta,
                "proyecto": proyecto,
                "fecha": fecha,
                "lugar": lugar,
                "hora_inicio": hora_inicio,
                "hora_fin": hora_fin,
                "resumen_ejecutivo": resumen_ejecutivo,
            },
            editable,
            resumen_reunion,
        )

        st.text_area("Vista previa acta completa (markdown)", value=md_full, height=260)

        c6, c7 = st.columns(2)
        with c6:
            st.download_button(
                "⬇️ Descargar acta completa (.md)",
                data=md_full.encode("utf-8"),
                file_name=f"acta_{no_acta or 'hoy'}.md",
                mime="text/markdown",
            )
        with c7:
            st.download_button(
                "⬇️ Descargar acta completa (.docx)",
                data=to_docx_from_md(md_full),
                file_name=f"acta_{no_acta or 'hoy'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        st.success("✅ Botón de 'acta de hoy' listo: completa campos y descarga .docx")
    else:
        st.warning("Primero carga y procesa un Excel en el Tab 1 para habilitar acta completa.")

st.divider()
st.caption("v3: semáforo por estado + acta completa en 1 clic")
