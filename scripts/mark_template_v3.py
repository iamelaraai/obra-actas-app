from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

SRC = Path('/Users/elara/.openclaw/workspace/obra_actas_app/templates/template_acta_oficial_v2.docx')
OUT = Path('/Users/elara/.openclaw/workspace/obra_actas_app/templates/template_acta_oficial_v3_marked.docx')


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def contains(p: Paragraph, needle: str) -> bool:
    return needle.lower() in (p.text or '').lower()


doc = Document(SRC)

def iter_all_paragraphs(document: Document):
    for p in document.paragraphs:
        yield p
    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

all_paras = list(iter_all_paragraphs(doc))

# Ensure global placeholders exist near header area
for p in all_paras:
    if contains(p, 'Acta del Comité de Obra') and '{{acta_no}}' not in p.text:
        p.text = 'Acta del Comité de Obra No. {{acta_no}}'
    if contains(p, 'Fecha:') and '{{fecha_larga}}' not in p.text:
        p.text = 'Fecha: {{fecha_larga}}'
    if contains(p, 'Lugar:') and '{{lugar}}' not in p.text:
        p.text = 'Lugar: {{lugar}}'
    if contains(p, 'Hora de inicio') and '{{hora_inicio}}' not in p.text:
        p.text = 'Hora de inicio: {{hora_inicio}}    Hora de finalización: {{hora_fin}}'
    if contains(p, 'Asistentes:') and '{{asistentes_total}}' not in p.text:
        p.text = 'Asistentes: {{asistentes_total}}'

# Insert section placeholders after known section titles (if found)
section_markers = [
    ('Por la SIF', '{{asistentes_sif}}'),
    ('Por la EDU', '{{asistentes_edu}}'),
    ('Por la interventoría', '{{asistentes_interventoria}}'),
    ('Por la empresa contratista', '{{asistentes_contratista}}'),
    ('Comité Técnico', '{{resumen_comite_tecnico}}'),
    ('Actividades ejecutadas durante la semana', '{{tabla_actividades}}'),
    ('Compromisos, Comentarios y observaciones', '{{tabla_compromisos}}'),
]

inserted = set()
for p in all_paras:
    for trigger, marker in section_markers:
        key = f'{trigger}|{marker}'
        if key in inserted:
            continue
        if contains(p, trigger):
            insert_paragraph_after(p, marker)
            inserted.add(key)

# If template is intentionally short, append canonical placeholder sections at end
required_markers = [
    '{{asistentes_sif}}',
    '{{asistentes_edu}}',
    '{{asistentes_interventoria}}',
    '{{asistentes_contratista}}',
    '{{resumen_comite_tecnico}}',
    '{{tabla_actividades}}',
    '{{tabla_compromisos}}',
]

existing_text = '\n'.join([p.text for p in all_paras if p.text])
missing = [m for m in required_markers if m not in existing_text]

if missing:
    doc.add_paragraph('')
    doc.add_heading('SECCIONES AUTOLLENABLES', level=1)

    doc.add_heading('Asistentes por entidad', level=2)
    doc.add_paragraph('{{asistentes_sif}}')
    doc.add_paragraph('{{asistentes_edu}}')
    doc.add_paragraph('{{asistentes_interventoria}}')
    doc.add_paragraph('{{asistentes_contratista}}')

    doc.add_heading('Comité Técnico', level=2)
    doc.add_paragraph('{{resumen_comite_tecnico}}')

    doc.add_heading('Actividades ejecutadas/proyectadas', level=2)
    doc.add_paragraph('{{tabla_actividades}}')

    doc.add_heading('Compromisos, comentarios y observaciones', level=2)
    doc.add_paragraph('{{tabla_compromisos}}')

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f'Created: {OUT}')
print(f'Inserted markers after existing sections: {len(inserted)}')
print(f'Missing markers appended as canonical block: {len(missing)}')
