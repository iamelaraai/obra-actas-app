import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def iter_paragraphs(doc):
    # body
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p

    # headers/footers
    for section in doc.sections:
        for container in [section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer]:
            if container is None:
                continue
            for p in container.paragraphs:
                yield p
            for t in container.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield p


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def style_cell_text(cell, bold=False, color_hex='000000', size_pt=9, align='left'):
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color_hex)


def replace_all_text(doc, mapping):
    for p in iter_paragraphs(doc):
        txt = p.text or ""
        new_txt = txt
        for k, v in mapping.items():
            if k in new_txt:
                new_txt = new_txt.replace(k, str(v))
        if new_txt != txt:
            p.text = new_txt


def find_marker_paragraph(doc, marker):
    for p in doc.paragraphs:
        if marker in (p.text or ""):
            return p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if marker in (p.text or ""):
                        return p
    return None


def make_table(doc, headers, rows, status_col_idx=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header style
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1F3864')
        style_cell_text(cell, bold=True, color_hex='FFFFFF', size_pt=9, align='center')

    if not rows:
        rows = [["Sin registros"] + [""] * (len(headers) - 1)]

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i in range(len(headers)):
            val = str(row[i]) if i < len(row) else ""
            cells[i].text = val

            # zebra rows
            if ridx % 2 == 1:
                set_cell_shading(cells[i], 'F2F2F2')

            # status semantic color
            if status_col_idx is not None and i == status_col_idx:
                v = val.strip().lower()
                if v == 'cumplido':
                    set_cell_shading(cells[i], 'D4EDDA')
                elif v in ('en proceso', 'cumplido parcialmente', 'pendiente', 'pendiente por definir'):
                    set_cell_shading(cells[i], 'FFF3CD')
                elif v == 'no cumplido':
                    set_cell_shading(cells[i], 'F8D7DA')

            style_cell_text(cells[i], bold=False, color_hex='111111', size_pt=9, align='left')

    return table


def insert_table_after_paragraph(doc, paragraph, headers, rows, status_col_idx=None):
    table = make_table(doc, headers, rows, status_col_idx=status_col_idx)
    try:
        paragraph._p.addnext(table._tbl)
    except Exception:
        # fallback: keep table at end if Word structure is unusual
        pass
    # add spacing paragraph after each inserted table
    doc.add_paragraph('')
    return table


def place_table_or_append(doc, marker, title, headers, rows, status_col_idx=None):
    p = find_marker_paragraph(doc, marker)
    if p is not None:
        p.text = title
        # style heading
        if p.runs:
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor.from_string('1F3864')
        insert_table_after_paragraph(doc, p, headers, rows, status_col_idx=status_col_idx)
        doc.add_paragraph('')
    else:
        h = doc.add_heading(title, level=2)
        if h.runs:
            for r in h.runs:
                r.font.color.rgb = RGBColor.from_string('1F3864')
        make_table(doc, headers, rows, status_col_idx=status_col_idx)
        doc.add_paragraph('')


def place_grouped_commitments(doc, marker, rows):
    headers = ["COMPROMISOS", "COMPONENTE", "RESPONSABLE", "FECHA/COMENTARIOS", "OBSERVACIONES"]
    groups = ["EDU", "Interventoría", "Contratista"]

    p = find_marker_paragraph(doc, marker)
    anchor = p
    if p is not None:
        p.text = "Compromisos, comentarios y observaciones"
    else:
        doc.add_heading("Compromisos, comentarios y observaciones", level=2)
        anchor = doc.paragraphs[-1]

    current_anchor = anchor
    for g in groups:
        sub = [r for r in rows if (r.get("responsable") or r.get("actor") or "").strip().lower() == g.lower()]
        if not sub:
            continue

        doc.add_paragraph('')
        # heading paragraph
        hp = doc.add_paragraph(f"Compromisos {g}:")
        try:
            current_anchor._p.addnext(hp._p)
        except Exception:
            pass

        data = []
        for r in sub:
            fc = "\n".join([x for x in [r.get("fechaLimite", ""), r.get("estado", "")] if x])
            data.append([
                r.get("compromiso", ""),
                r.get("componente", ""),
                r.get("responsable", "") or r.get("actor", ""),
                fc,
                r.get("observacion", ""),
            ])

        t = insert_table_after_paragraph(doc, hp, headers, data)
        current_anchor = hp


def main():
    if len(sys.argv) != 4:
        print("Usage: generate_acta_docx.py <template.docx> <payload.json> <output.docx>")
        sys.exit(1)

    template_path = Path(sys.argv[1])
    payload_path = Path(sys.argv[2])
    out_path = Path(sys.argv[3])

    payload = json.loads(payload_path.read_text(encoding="utf-8"))

    doc = Document(str(template_path))

    asistentes = payload.get("asistentes", {})
    total_asist = sum(len(asistentes.get(k, [])) for k in ["sif", "edu", "interventoria", "contratista"])

    mapping = {
        "{{objeto_proyecto}}": payload.get("meta", {}).get("objeto_proyecto", ""),
        "{{acta_no}}": payload.get("meta", {}).get("acta_no", ""),
        "{{fecha_larga}}": payload.get("meta", {}).get("fecha_larga", ""),
        "{{lugar}}": payload.get("meta", {}).get("lugar", ""),
        "{{hora_inicio}}": payload.get("meta", {}).get("hora_inicio", ""),
        "{{hora_fin}}": payload.get("meta", {}).get("hora_fin", ""),
        "{{asistentes_total}}": str(total_asist),
        "{{resumen_comite_tecnico}}": payload.get("resumenReunion", ""),
        "{{pagina}}": payload.get("meta", {}).get("pagina", "1"),
        "{{página}}": payload.get("meta", {}).get("pagina", "1"),
        "{{página)}}": payload.get("meta", {}).get("pagina", "1"),
    }

    replace_all_text(doc, mapping)

    def det(key):
        arr = asistentes.get(f"{key}_det")
        if arr:
            return [[x.get("nombre", ""), x.get("cargo", "")] for x in arr]
        return [[n, ""] for n in asistentes.get(key, [])]

    # 1) Asistentes al inicio en formato tabla por entidad
    place_table_or_append(doc, "{{asistentes_sif}}", "Por la SIF", ["NOMBRE", "CARGO"], det("sif"))
    place_table_or_append(doc, "{{asistentes_edu}}", "Por la EDU", ["NOMBRE", "CARGO"], det("edu"))
    place_table_or_append(doc, "{{asistentes_interventoria}}", "Por la Interventoría", ["NOMBRE", "CARGO"], det("interventoria"))
    place_table_or_append(doc, "{{asistentes_contratista}}", "Por la empresa contratista", ["NOMBRE", "CARGO"], det("contratista"))

    rows = payload.get("rows", [])

    # 2) Tabla de actividades general (sin columnas ACTOR y COMPONENTE)
    data_rows = [
        [
            r.get("compromiso", ""),
            r.get("responsable", "") or r.get("actor", ""),
            r.get("fechaLimite", ""),
            r.get("estado", ""),
            r.get("observacion", ""),
        ]
        for r in rows
    ]
    headers_actividad = ["COMPROMISO", "RESPONSABLE", "FECHA", "ESTADO", "OBSERVACIÓN"]
    place_table_or_append(doc, "{{tabla_actividades}}", "Tabla de actividades", headers_actividad, data_rows, status_col_idx=3)

    # 3) Compromisos por entidad (formato solicitado)
    place_grouped_commitments(doc, "{{tabla_compromisos}}", rows)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    main()
