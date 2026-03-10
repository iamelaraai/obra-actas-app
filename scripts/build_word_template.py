from pathlib import Path
from docx import Document

SRC = Path('/Users/elara/.openclaw/media/inbound/ACTA_18_Comité_03-03-2026_Parque_Primavera_Norte_V0---343f960e-c722-48b2-9feb-68b5aaf303e9.docx')
OUT = Path('/Users/elara/.openclaw/workspace/obra_actas_app/templates/template_acta_oficial_v1.docx')

replacements = {
    'Acta del Comité de Obra No. 18': 'Acta del Comité de Obra No. {{acta_no}}',
    'Fecha:  03 de marzo de 2026': 'Fecha:  {{fecha_larga}}',
    'Lugar:  Campamento de obra': 'Lugar:  {{lugar}}',
    'Hora de inicio:  9:30 AM       \tHora de finalización:  11:58 AM': 'Hora de inicio:  {{hora_inicio}}       \tHora de finalización:  {{hora_fin}}',
    'Asistentes: 31': 'Asistentes: {{asistentes_total}}',
    '“CONSTRUCCIÓN DEL ESPACIO PUBLICO Y EQUIPAMIENTOS DEL PARQUE PRIMAVERA NORTE UBICADO EN EL DISTRITO DE CIENCIA, TECNOLOGÍA E INNOVACIÓN DE MEDELLÍN”': '“{{objeto_proyecto}}”',
}

# helper: replace text inside paragraph preserving first run style (simple approach)
def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    # collapse to one run for deterministic replacement
    txt = paragraph.text.replace(old, new)
    for _ in range(len(paragraph.runs)-1, -1, -1):
        paragraph._p.remove(paragraph.runs[_]._r)
    paragraph.add_run(txt)
    return True


doc = Document(SRC)

for p in doc.paragraphs:
    for old, new in replacements.items():
        if old in p.text:
            replace_in_paragraph(p, old, new)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f'Created {OUT}')
